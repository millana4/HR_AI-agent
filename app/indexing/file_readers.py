"""
Скачивание файлов из CDN и извлечение текста.

Поддерживаемые форматы: PDF, DOCX, TXT.
Формат определяется по расширению в URL.

Использование:
    text = await download_and_extract("https://cdn.example.com/doc.pdf")
"""
import io
from urllib.parse import urlparse

import httpx
from docx import Document as DocxDocument
from pypdf import PdfReader

from app.core.exceptions import RepositoryError
from app.core.logging import get_logger


logger = get_logger(__name__)


# Таймаут скачивания файла. PDF до 100 страниц весит ~5-20 МБ —
# 60 секунд должно хватать с большим запасом для CDN.
DOWNLOAD_TIMEOUT = 60.0


async def download_and_extract(
    url: str,
    correlation_id: str = "-",
) -> str:
    """
    Скачать файл из URL и извлечь текст.

    Args:
        url: HTTPS-ссылка на файл в CDN
        correlation_id: для трассировки

    Returns:
        Извлечённый текст. Может быть пустой строкой, если файл — PDF-скан
        без текстового слоя.

    Raises:
        RepositoryError: если скачивание не удалось, формат не поддерживается,
                         или файл битый.
    """
    file_format = _detect_format(url)
    logger.debug(
        f"Downloading {file_format} from {url}",
        extra={"correlation_id": correlation_id},
    )

    content = await _download(url, correlation_id=correlation_id)

    if file_format == "pdf":
        return _extract_pdf(content, url=url, correlation_id=correlation_id)
    elif file_format == "docx":
        return _extract_docx(content, url=url, correlation_id=correlation_id)
    elif file_format == "txt":
        return _extract_txt(content, url=url, correlation_id=correlation_id)
    else:
        # _detect_format уже отфильтровал, но для строгости типов
        raise RepositoryError(f"Unsupported format: {file_format}")


def _detect_format(url: str) -> str:
    """
    Определить формат файла по расширению в URL.

    Поддерживает .pdf, .docx, .txt. Регистр не важен.
    """
    path = urlparse(url).path.lower()
    if path.endswith(".pdf"):
        return "pdf"
    elif path.endswith(".docx"):
        return "docx"
    elif path.endswith(".doc"):
        # .doc (старый формат) python-docx читает с переменным успехом.
        # Бланки Вотоня сохранены как .doc — попробуем через docx.
        # Если не справится — отдельно решим (можно конвертировать вручную в .docx).
        return "docx"
    elif path.endswith(".txt"):
        return "txt"
    else:
        raise RepositoryError(
            f"Cannot detect format from URL: {url}. "
            f"Supported extensions: .pdf, .docx, .doc, .txt"
        )


async def _download(url: str, correlation_id: str) -> bytes:
    """Скачать файл, вернуть его байты."""
    async with httpx.AsyncClient(
        timeout=DOWNLOAD_TIMEOUT,
        trust_env=False,
        follow_redirects=True,
    ) as client:
        try:
            response = await client.get(url)
            response.raise_for_status()
        except httpx.HTTPStatusError as exc:
            raise RepositoryError(
                f"Failed to download {url}: HTTP {exc.response.status_code}"
            ) from exc
        except httpx.HTTPError as exc:
            raise RepositoryError(f"Failed to download {url}: {exc}") from exc

    logger.debug(
        f"Downloaded {len(response.content)} bytes from {url}",
        extra={"correlation_id": correlation_id},
    )
    return response.content


def _extract_pdf(content: bytes, url: str, correlation_id: str) -> str:
    """
    Извлечь текст из PDF.

    Если PDF — скан без текстового слоя, у каждой страницы будет пустой текст.
    В этом случае возвращаем "" и логируем warning — индексатор пропустит документ.
    """
    try:
        reader = PdfReader(io.BytesIO(content))
        pages_text = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise RepositoryError(f"Failed to parse PDF {url}: {exc}") from exc

    text = "\n\n".join(p.strip() for p in pages_text if p.strip())

    if not text:
        logger.warning(
            f"PDF has no extractable text (likely scan): {url}",
            extra={"correlation_id": correlation_id},
        )

    return text


def _extract_docx(content: bytes, url: str, correlation_id: str) -> str:
    """
    Извлечь текст из DOCX.

    Берём текст всех параграфов + ячеек таблиц.
    """
    try:
        doc = DocxDocument(io.BytesIO(content))
    except Exception as exc:
        raise RepositoryError(f"Failed to parse DOCX {url}: {exc}") from exc

    parts: list[str] = []
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return "\n\n".join(parts)


def _extract_txt(content: bytes, url: str, correlation_id: str) -> str:
    """Извлечь текст из TXT. Пробуем UTF-8, потом cp1251 (legacy Windows)."""
    try:
        return content.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return content.decode("cp1251")
        except UnicodeDecodeError as exc:
            raise RepositoryError(f"Failed to decode TXT {url}: {exc}") from exc